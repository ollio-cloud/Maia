#!/usr/bin/env python3
"""
Intune Policy Configuration Reporter
Extracts Intune policies starting with "Orro" and generates a configuration report.

Prerequisites:
- Azure AD App Registration with Microsoft Graph API permissions:
  - DeviceManagementConfiguration.Read.All
  - DeviceManagementApps.Read.All
  - DeviceManagementManagedDevices.Read.All
- python-docx library for Word document generation
- msal library for authentication

Usage:
    python intune_policy_reporter.py --tenant-id YOUR_TENANT_ID --client-id YOUR_CLIENT_ID --client-secret YOUR_SECRET
"""

import argparse
import sys
from datetime import datetime
from typing import List, Dict, Any
import json

try:
    import msal
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import requests
except ImportError as e:
    print(f"Error: Missing required library: {e}")
    print("Please install required packages:")
    print("  pip install msal python-docx requests")
    sys.exit(1)


class IntuneGraphClient:
    """Client for Microsoft Graph API to interact with Intune."""

    def __init__(self, tenant_id: str, client_id: str, client_secret: str):
        self.tenant_id = tenant_id
        self.client_id = client_id
        self.client_secret = client_secret
        self.access_token = None
        self.base_url = "https://graph.microsoft.com/v1.0"

    def authenticate(self) -> bool:
        """Authenticate using MSAL and get access token."""
        authority = f"https://login.microsoftonline.com/{self.tenant_id}"
        app = msal.ConfidentialClientApplication(
            self.client_id,
            authority=authority,
            client_credential=self.client_secret
        )

        scopes = ["https://graph.microsoft.com/.default"]
        result = app.acquire_token_for_client(scopes=scopes)

        if "access_token" in result:
            self.access_token = result["access_token"]
            print("✓ Successfully authenticated to Microsoft Graph API")
            return True
        else:
            print(f"✗ Authentication failed: {result.get('error_description', 'Unknown error')}")
            return False

    def _make_request(self, endpoint: str) -> Dict[str, Any]:
        """Make authenticated request to Microsoft Graph API."""
        if not self.access_token:
            raise Exception("Not authenticated. Call authenticate() first.")

        headers = {
            "Authorization": f"Bearer {self.access_token}",
            "Content-Type": "application/json"
        }

        url = f"{self.base_url}{endpoint}"
        response = requests.get(url, headers=headers)

        if response.status_code == 200:
            return response.json()
        else:
            print(f"Warning: Request to {endpoint} failed with status {response.status_code}")
            return {"value": []}

    def get_device_configurations(self) -> List[Dict[str, Any]]:
        """Get all device configuration policies."""
        print("Fetching device configuration policies...")
        result = self._make_request("/deviceManagement/deviceConfigurations")
        policies = result.get("value", [])
        print(f"  Found {len(policies)} device configuration policies")
        return policies

    def get_device_compliance_policies(self) -> List[Dict[str, Any]]:
        """Get all device compliance policies."""
        print("Fetching device compliance policies...")
        result = self._make_request("/deviceManagement/deviceCompliancePolicies")
        policies = result.get("value", [])
        print(f"  Found {len(policies)} compliance policies")
        return policies

    def get_app_protection_policies(self) -> List[Dict[str, Any]]:
        """Get all app protection policies (iOS and Android)."""
        print("Fetching app protection policies...")

        # iOS policies
        ios_result = self._make_request("/deviceAppManagement/iosManagedAppProtections")
        ios_policies = ios_result.get("value", [])

        # Android policies
        android_result = self._make_request("/deviceAppManagement/androidManagedAppProtections")
        android_policies = android_result.get("value", [])

        # Windows policies
        windows_result = self._make_request("/deviceAppManagement/windowsManagedAppProtections")
        windows_policies = windows_result.get("value", [])

        all_policies = ios_policies + android_policies + windows_policies
        print(f"  Found {len(all_policies)} app protection policies")
        return all_policies

    def get_configuration_profiles(self) -> List[Dict[str, Any]]:
        """Get Settings Catalog and Administrative Templates."""
        print("Fetching configuration profiles...")
        result = self._make_request("/deviceManagement/configurationPolicies")
        policies = result.get("value", [])
        print(f"  Found {len(policies)} configuration profiles")
        return policies

    def get_policy_details(self, policy_id: str, policy_type: str) -> Dict[str, Any]:
        """Get detailed settings for a specific policy."""
        endpoint_map = {
            "deviceConfiguration": f"/deviceManagement/deviceConfigurations/{policy_id}",
            "compliancePolicy": f"/deviceManagement/deviceCompliancePolicies/{policy_id}",
            "configurationPolicy": f"/deviceManagement/configurationPolicies/{policy_id}",
        }

        endpoint = endpoint_map.get(policy_type)
        if endpoint:
            return self._make_request(endpoint)
        return {}


class IntunePolicyReporter:
    """Generate Word document reports from Intune policies."""

    def __init__(self, filter_prefix: str = "Orro"):
        self.filter_prefix = filter_prefix
        self.doc = Document()

    def filter_policies(self, policies: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Filter policies that start with the specified prefix."""
        filtered = [p for p in policies if p.get("displayName", "").startswith(self.filter_prefix)]
        return filtered

    def create_report(self, all_policies: Dict[str, List[Dict[str, Any]]], output_path: str):
        """Create a comprehensive Word document report."""

        # Title page
        title = self.doc.add_heading('Intune Configuration Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph(f'Organization: {self.filter_prefix}')
        self.doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        self.doc.add_paragraph(f'Filter: Policies starting with "{self.filter_prefix}"')
        self.doc.add_paragraph()

        # Executive Summary
        self.doc.add_heading('Executive Summary', 1)

        summary_table = self.doc.add_table(rows=1, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = 'Policy Type'
        header_cells[1].text = 'Count'

        for policy_type, policies in all_policies.items():
            filtered = self.filter_policies(policies)
            if filtered:
                row = summary_table.add_row().cells
                row[0].text = policy_type.replace('_', ' ').title()
                row[1].text = str(len(filtered))

        self.doc.add_paragraph()

        # Detailed Policy Sections
        for policy_type, policies in all_policies.items():
            filtered = self.filter_policies(policies)

            if not filtered:
                continue

            self.doc.add_page_break()
            self.doc.add_heading(policy_type.replace('_', ' ').title(), 1)

            for policy in filtered:
                self._add_policy_section(policy, policy_type)

        # Save document
        self.doc.save(output_path)
        print(f"\n✓ Report generated: {output_path}")

    def _add_policy_section(self, policy: Dict[str, Any], policy_type: str):
        """Add a section for a single policy."""

        # Policy name as heading
        policy_name = policy.get('displayName', 'Unnamed Policy')
        self.doc.add_heading(policy_name, 2)

        # Basic information table
        info_table = self.doc.add_table(rows=0, cols=2)
        info_table.style = 'Light List Accent 1'

        # Add key fields
        fields = [
            ('Display Name', policy.get('displayName', 'N/A')),
            ('Description', policy.get('description', 'No description')),
            ('Platform', policy.get('@odata.type', 'Unknown').split('.')[-1]),
            ('Created', policy.get('createdDateTime', 'N/A')),
            ('Last Modified', policy.get('lastModifiedDateTime', 'N/A')),
            ('ID', policy.get('id', 'N/A')),
        ]

        for field_name, field_value in fields:
            row = info_table.add_row().cells
            row[0].text = field_name
            row[0].paragraphs[0].runs[0].bold = True
            row[1].text = str(field_value)

        # Add assignments if available
        if 'assignments' in policy:
            self.doc.add_paragraph()
            self.doc.add_heading('Assignments', 3)
            assignments = policy.get('assignments', [])
            if assignments:
                for assignment in assignments:
                    target = assignment.get('target', {})
                    self.doc.add_paragraph(
                        f"• {target.get('@odata.type', 'Unknown').split('.')[-1]}: {target.get('groupId', 'All Users/Devices')}",
                        style='List Bullet'
                    )
            else:
                self.doc.add_paragraph('No assignments configured')

        # Add settings summary
        self.doc.add_paragraph()
        self.doc.add_heading('Configuration Settings', 3)

        # Filter out metadata fields and show actual settings
        settings = {k: v for k, v in policy.items()
                   if not k.startswith('@') and k not in ['id', 'displayName', 'description',
                                                           'createdDateTime', 'lastModifiedDateTime',
                                                           'version', 'assignments']}

        if settings:
            settings_text = json.dumps(settings, indent=2, default=str)
            # Limit to first 2000 chars to avoid huge reports
            if len(settings_text) > 2000:
                settings_text = settings_text[:2000] + "\n... (truncated)"

            p = self.doc.add_paragraph(settings_text)
            p.style = 'Normal'
            # Use monospace font
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        else:
            self.doc.add_paragraph('No additional settings found')

        self.doc.add_paragraph()


def main():
    parser = argparse.ArgumentParser(
        description='Generate Intune configuration report for policies starting with "Orro"'
    )
    parser.add_argument('--tenant-id', required=True, help='Azure AD Tenant ID')
    parser.add_argument('--client-id', required=True, help='Azure AD App Client ID')
    parser.add_argument('--client-secret', required=True, help='Azure AD App Client Secret')
    parser.add_argument('--prefix', default='Orro', help='Policy name prefix to filter (default: Orro)')
    parser.add_argument('--output', default='Intune_Orro_Configuration_Report.docx',
                       help='Output file path (default: Intune_Orro_Configuration_Report.docx)')

    args = parser.parse_args()

    print("=" * 60)
    print("Intune Policy Configuration Reporter")
    print("=" * 60)
    print()

    # Authenticate and connect to Graph API
    client = IntuneGraphClient(args.tenant_id, args.client_id, args.client_secret)

    if not client.authenticate():
        print("\n✗ Authentication failed. Please check your credentials.")
        sys.exit(1)

    print()

    # Fetch all policies
    all_policies = {
        'device_configurations': client.get_device_configurations(),
        'compliance_policies': client.get_device_compliance_policies(),
        'app_protection_policies': client.get_app_protection_policies(),
        'configuration_profiles': client.get_configuration_profiles(),
    }

    print()

    # Generate report
    reporter = IntunePolicyReporter(filter_prefix=args.prefix)

    # Count filtered policies
    total_filtered = sum(len(reporter.filter_policies(policies))
                        for policies in all_policies.values())

    print(f"Generating report for {total_filtered} policies starting with '{args.prefix}'...")
    reporter.create_report(all_policies, args.output)

    print()
    print("=" * 60)
    print("Report generation complete!")
    print("=" * 60)


if __name__ == "__main__":
    main()
