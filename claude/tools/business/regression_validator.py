#!/usr/bin/env python3
"""
Regression validation framework
Creates baselines, runs before/after comparisons, validates changes don't break existing functionality
"""

import json
import hashlib
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional
from docx import Document
from md2docx_v2 import convert_markdown_to_docx

class RegressionValidator:
    """Validates changes don't break existing functionality"""

    def __init__(self, test_dir: Path = None):
        self.test_dir = test_dir or Path("regression_tests")
        self.baseline_dir = self.test_dir / "baselines"
        self.temp_dir = self.test_dir / "temp"
        self.reports_dir = self.test_dir / "reports"

        # Ensure directories exist
        for dir_path in [self.test_dir, self.baseline_dir, self.temp_dir, self.reports_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)

    def create_baseline(self, test_name: str, md_content: str, modes: List[str] = None) -> Dict[str, Any]:
        """Create baseline for a test case"""
        modes = modes or ["styled", "ats", "readable"]

        baseline = {
            "test_name": test_name,
            "created": datetime.now().isoformat(),
            "md_content": md_content,
            "md_hash": hashlib.md5(md_content.encode(), usedforsecurity=False).hexdigest(),
            "modes": {}
        }

        # Create baseline for each mode
        for mode in modes:
            baseline["modes"][mode] = self._create_mode_baseline(test_name, md_content, mode)

        # Save baseline
        baseline_file = self.baseline_dir / f"{test_name}.json"
        with open(baseline_file, 'w', encoding='utf-8') as f:
            json.dump(baseline, f, indent=2)

        return baseline

    def _create_mode_baseline(self, test_name: str, md_content: str, mode: str) -> Dict[str, Any]:
        """Create baseline for specific mode"""
        # Create temp files
        md_file = self.temp_dir / f"{test_name}_{mode}.md"
        docx_file = self.temp_dir / f"{test_name}_{mode}.docx"

        md_file.write_text(md_content, encoding='utf-8')

        # Convert
        convert_markdown_to_docx(md_file, docx_file, mode)

        # Analyze document
        doc = Document(docx_file)
        analysis = self._analyze_document_structure(doc)

        # Save baseline document
        baseline_docx = self.baseline_dir / f"{test_name}_{mode}_baseline.docx"
        shutil.copy2(docx_file, baseline_docx)

        return analysis

    def validate_against_baseline(self, test_name: str, md_content: str = None, modes: List[str] = None) -> Dict[str, Any]:
        """Validate current conversion against baseline"""
        baseline_file = self.baseline_dir / f"{test_name}.json"

        if not baseline_file.exists():
            return {"error": f"No baseline found for {test_name}"}

        # Load baseline
        with open(baseline_file, 'r', encoding='utf-8') as f:
            baseline = json.load(f)

        # Use baseline markdown if not provided
        if md_content is None:
            md_content = baseline["md_content"]

        modes = modes or list(baseline["modes"].keys())

        # Run current conversion
        current_results = {}
        for mode in modes:
            current_results[mode] = self._create_mode_baseline(f"{test_name}_current", md_content, mode)

        # Compare results
        comparison = self._compare_results(baseline["modes"], current_results)

        return {
            "test_name": test_name,
            "timestamp": datetime.now().isoformat(),
            "baseline_date": baseline["created"],
            "modes_tested": modes,
            "comparison": comparison,
            "overall_status": "PASS" if all(m["status"] == "PASS" for m in comparison.values()) else "FAIL"
        }

    def _compare_results(self, baseline: Dict[str, Any], current: Dict[str, Any]) -> Dict[str, Any]:
        """Compare baseline and current results"""
        comparison = {}

        for mode in baseline.keys():
            if mode not in current:
                comparison[mode] = {"status": "SKIP", "reason": f"Mode {mode} not tested"}
                continue

            mode_comparison = {
                "status": "PASS",
                "differences": [],
                "metrics": {}
            }

            baseline_data = baseline[mode]
            current_data = current[mode]

            # Compare key metrics
            metrics = ["total_paragraphs", "section_count", "font_count", "style_count"]
            for metric in metrics:
                baseline_val = baseline_data.get(metric, 0)
                current_val = current_data.get(metric, 0)
                mode_comparison["metrics"][metric] = {
                    "baseline": baseline_val,
                    "current": current_val,
                    "change": current_val - baseline_val
                }

                # Flag significant changes
                if abs(current_val - baseline_val) > 1:
                    mode_comparison["differences"].append({
                        "type": "metric_change",
                        "metric": metric,
                        "baseline": baseline_val,
                        "current": current_val
                    })

            # Compare structure
            self._compare_structure(baseline_data, current_data, mode_comparison)

            # Determine status
            if mode_comparison["differences"]:
                mode_comparison["status"] = "FAIL"

            comparison[mode] = mode_comparison

        return comparison

    def _compare_structure(self, baseline: Dict[str, Any], current: Dict[str, Any], comparison: Dict[str, Any]):
        """Compare document structure"""
        # Compare sections
        baseline_sections = set(s["title"] for s in baseline.get("sections", []))
        current_sections = set(s["title"] for s in current.get("sections", []))

        missing_sections = baseline_sections - current_sections
        extra_sections = current_sections - baseline_sections

        if missing_sections:
            comparison["differences"].append({
                "type": "missing_sections",
                "sections": list(missing_sections)
            })

        if extra_sections:
            comparison["differences"].append({
                "type": "extra_sections",
                "sections": list(extra_sections)
            })

        # Compare fonts
        baseline_fonts = set(baseline.get("fonts", []))
        current_fonts = set(current.get("fonts", []))

        if baseline_fonts != current_fonts:
            comparison["differences"].append({
                "type": "font_changes",
                "baseline_fonts": list(baseline_fonts),
                "current_fonts": list(current_fonts)
            })

    def _analyze_document_structure(self, doc: Document) -> Dict[str, Any]:
        """Analyze document structure for comparison"""
        analysis = {
            "total_paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
            "sections": [],
            "fonts": set(),
            "styles": set(),
            "section_count": 0,
            "font_count": 0,
            "style_count": 0
        }

        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            analysis["styles"].add(para.style.name)

            # Check for section headers
            if text in ["Professional Summary", "Key Achievements", "Professional Experience", "Certifications", "Referees"]:
                if current_section:
                    analysis["sections"].append(current_section)
                current_section = {
                    "title": text,
                    "content_count": 0
                }
                analysis["section_count"] += 1
            elif current_section:
                current_section["content_count"] += 1

            # Collect fonts
            for run in para.runs:
                if run.font.name:
                    analysis["fonts"].add(run.font.name)

        if current_section:
            analysis["sections"].append(current_section)

        # Convert sets to lists for JSON serialization
        analysis["fonts"] = list(analysis["fonts"])
        analysis["styles"] = list(analysis["styles"])
        analysis["font_count"] = len(analysis["fonts"])
        analysis["style_count"] = len(analysis["styles"])

        return analysis

    def run_all_baselines(self) -> Dict[str, Any]:
        """Run regression tests against all baselines"""
        results = {
            "timestamp": datetime.now().isoformat(),
            "tests": {},
            "summary": {"total": 0, "passed": 0, "failed": 0}
        }

        # Find all baseline files
        baseline_files = list(self.baseline_dir.glob("*.json"))

        for baseline_file in baseline_files:
            test_name = baseline_file.stem
            if test_name.endswith("_current"):
                continue  # Skip temp files

            print(f"Running regression test: {test_name}")
            test_result = self.validate_against_baseline(test_name)
            results["tests"][test_name] = test_result

            results["summary"]["total"] += 1
            if test_result.get("overall_status") == "PASS":
                results["summary"]["passed"] += 1
            else:
                results["summary"]["failed"] += 1

        # Save report
        report_file = self.reports_dir / f"regression_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2)

        return results

    def create_standard_baselines(self):
        """Create baselines for standard test cases"""
        test_cases = {
            "minimal_cv": """# John Doe
**Software Engineer**

Email: john@example.com

## Professional Summary

Experienced software engineer.
""",

            "full_cv": """# Jane Smith
**Product Manager**

Location: London, UK
Mobile: +44 123 456 7890
Email: jane@example.com
https://linkedin.com/in/janesmith

## Professional Summary

Product manager with experience.

## Key Achievements

• **Led product launch** — Successful launch
• **Increased revenue** — 25% growth

## Professional Experience

**TechCorp** | London, UK
**Senior Product Manager** | Jan 2020 - Present

• **Managed roadmap** — Coordinated with teams

## Certifications

**Product Management:** Certified Product Manager

## Referees

Available upon request
""",

            "certifications_test": """# Test User
**Test Role**

Email: test@example.com

## Certifications

**Cloud Platforms:** AWS Certified | Azure Certified
**Programming:** Python Certified | Java Certified

## Referees

Available upon request
"""
        }

        print("Creating standard baselines...")
        for test_name, md_content in test_cases.items():
            print(f"  Creating baseline: {test_name}")
            self.create_baseline(test_name, md_content)

        print("Standard baselines created!")


def main():
    """Main function for regression validation"""
    import argparse

    parser = argparse.ArgumentParser(description="CV Converter Regression Validation")
    parser.add_argument("--create-baselines", action="store_true",
                       help="Create standard baselines")
    parser.add_argument("--run-regression", action="store_true",
                       help="Run regression tests against all baselines")
    parser.add_argument("--test-name", type=str,
                       help="Test specific baseline")

    args = parser.parse_args()

    validator = RegressionValidator()

    if args.create_baselines:
        validator.create_standard_baselines()

    elif args.run_regression:
        print("Running regression validation...")
        results = validator.run_all_baselines()

        print("\n" + "="*50)
        print("REGRESSION TEST RESULTS")
        print("="*50)
        print(f"Total tests: {results['summary']['total']}")
        print(f"Passed: {results['summary']['passed']}")
        print(f"Failed: {results['summary']['failed']}")
        print(f"Success rate: {(results['summary']['passed'] / results['summary']['total'] * 100):.1f}%")

        # Show failures
        for test_name, test_result in results["tests"].items():
            if test_result.get("overall_status") == "FAIL":
                print(f"\nFAILED: {test_name}")
                for mode, comparison in test_result["comparison"].items():
                    if comparison["status"] == "FAIL":
                        print(f"  Mode {mode}:")
                        for diff in comparison["differences"]:
                            print(f"    - {diff['type']}: {diff}")

    elif args.test_name:
        print(f"Testing baseline: {args.test_name}")
        result = validator.validate_against_baseline(args.test_name)
        print(f"Status: {result.get('overall_status', 'UNKNOWN')}")

    else:
        print("Use --create-baselines, --run-regression, or --test-name")


if __name__ == "__main__":
    main()
