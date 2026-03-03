#!/usr/bin/env python3
"""
Markdown to Word Document Converter with Strict Style Control
Supports both styled (human-readable) and ATS-optimized output modes
"""

import argparse
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import mistune
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class StyleConfig:
    """Configuration for document styling"""

    def __init__(self, mode: str = "styled"):
        self.mode = mode
        self.styles = self._get_style_config()

    def _get_style_config(self) -> Dict:
        """Return style configuration based on mode"""
        base = {
            "font": "Calibri",
            "margins": {
                "top": 0.39,
                "bottom": 0.59,
                "left": 0.79,
                "right": 0.88
            },
            "spacing": {
                "after_name": 6,
                "after_title": 6,
                "after_contact": 6,
                "after_summary": 6,  # Changed from 5pt to 6pt as requested
                "before_section": 12,
                "after_section_title": 6,
                "before_job": 0,      # No space before job entries
                "after_job_title": 0,    # No gap between job title and description
                "between_bullets": 2,
                "after_job_description": 0,  # No space after description
                "emphasis_spacing": 18.8  # For current role bullets
            }
        }

        if self.mode == "ats":
            # ATS-friendly: left-aligned, clear labels, simple formatting
            return {
                **base,
                "header_align": WD_ALIGN_PARAGRAPH.LEFT,
                "body_align": WD_ALIGN_PARAGRAPH.LEFT,
                "name_size": 16,
                "contact_size": 11,
                "h1_size": 14,
                "h2_size": 12,
                "body_size": 11,
                "bullet_size": 10,
                "description_size": 9,
                "company_size": 11,
                "role_size": 11,
                "use_labels": True,  # Add Email:, Phone: etc.
                "bullet_indent": 0.25,
                "line_spacing": 1.15
            }
        elif self.mode == "readable":
            # Human-readable: optimized for comfort and readability
            readable_spacing = base["spacing"].copy()
            readable_spacing.update({
                "before_job": 0,     # No space before job entries
                "after_job_title": 0, # No gap between title and description
                "between_bullets": 4,
                "after_job_description": 0,  # No space after description
                "after_summary": 8
            })
            return {
                **base,
                "spacing": readable_spacing,
                "header_align": WD_ALIGN_PARAGRAPH.LEFT,
                "body_align": WD_ALIGN_PARAGRAPH.LEFT,
                "name_size": 16,
                "contact_size": 11,
                "h1_size": 14,
                "h2_size": 12,
                "body_size": 11,
                "bullet_size": 10,
                "description_size": 9,
                "company_size": 11,
                "role_size": 11,
                "use_labels": False,
                "bullet_indent": 0.25,
                "line_spacing": 1.15
            }
        else:
            # Styled: matches reference document exactly
            return {
                **base,
                "header_align": WD_ALIGN_PARAGRAPH.RIGHT,
                "body_align": WD_ALIGN_PARAGRAPH.JUSTIFY,
                "name_size": 14,
                "contact_size": 10,
                "h1_size": 16,
                "h2_size": 12,
                "body_size": 10,
                "bullet_size": 10,
                "description_size": 9,
                "company_size": 11,
                "role_size": 11,
                "use_labels": False,
                "bullet_indent": 0.25,
                "line_spacing": 1.0
            }


class MarkdownParser:
    """Parse markdown and extract structured content"""

    def __init__(self, content: str):
        self.content = content
        self.lines = content.split('\n')
        self.parsed_data = self._parse_structure()

    def _extract_bold_text(self, text: str) -> Tuple[str, str]:
        """Extract text from between ** markers"""
        match = re.match(r'\*\*(.*?)\*\*(.*)$', text.strip())
        if match:
            return match.group(1), match.group(2).strip()
        return text.strip(), ""

    def _extract_italic_text(self, text: str) -> str:
        """Extract text from between single * markers (italic)"""
        match = re.match(r'\*(.*?)\*$', text.strip())
        if match:
            return match.group(1)
        return text.strip()

    def _clean_contact_text(self, text: str) -> str:
        """Remove markdown links and clean contact text"""
        # Remove markdown links [text](url) and replace with just text
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        return text.strip()

    def _parse_structure(self) -> Dict:
        """Extract CV structure from markdown"""
        data = {
            "name": "",
            "title": "",
            "contact": [],
            "sections": []
        }

        current_section = None
        current_subsection = None
        processed_lines = set()  # Track which lines we've already processed

        for i, line in enumerate(self.lines):
            line = line.strip()

            # Skip empty lines
            if not line:
                continue

            # Skip if we've already processed this line as part of a multi-line pattern
            if i in processed_lines:
                continue

            # Extract name (first H1)
            if line.startswith('# ') and not data["name"]:
                data["name"] = line[2:].strip()

            # Extract title (bold text immediately after name, before contact)
            elif line.startswith('**') and not data["title"] and data["name"] and not data["contact"]:
                title_text, _ = self._extract_bold_text(line)
                data["title"] = title_text

            # Extract contact info (single line with pipes or collect separate lines)
            elif ('|' in line and '@' in line) or (
                data["name"] and data["title"] and
                (line.startswith('Location:') or line.startswith('Nationality:') or
                 line.startswith('Mobile:') or line.startswith('Email:') or
                 '@' in line or line.startswith('http'))
            ):
                if not data["contact"]:
                    data["contact"] = []

                if '|' in line and '@' in line:
                    # Old format with pipes
                    cleaned_line = self._clean_contact_text(line)
                    data["contact"] = [item.strip() for item in cleaned_line.split('|')]
                else:
                    # New format with separate lines - collect this line
                    cleaned_line = self._clean_contact_text(line)
                    data["contact"].append(cleaned_line)

            # H2 sections
            elif line.startswith('## '):
                section_name = line[3:].strip()
                current_section = {
                    "title": section_name,
                    "content": [],
                    "subsections": []
                }
                data["sections"].append(current_section)
                current_subsection = None

            # Job entries (bold company name pattern OR additional role at same company)
            elif line.startswith('**') and current_section and '|' in line and current_section["title"] == "Professional Experience":
                parts = line.split('|')
                first_part, _ = self._extract_bold_text(parts[0])
                second_part = parts[1].strip() if len(parts) > 1 else ""

                # Check if this looks like a date range (additional role) or location (new company)
                is_date_range = any(date_indicator in second_part.lower() for date_indicator in
                                   ['jan', 'feb', 'mar', 'apr', 'may', 'jun',
                                    'jul', 'aug', 'sep', 'oct', 'nov', 'dec',
                                    '20', '19', '—', '-', 'present'])

                if is_date_range:
                    # This is a job title line (either additional role or first role)
                    if current_subsection and current_subsection.get("company"):
                        # Additional role at same company
                        new_role = {
                            "company": current_subsection["company"],
                            "location": current_subsection["location"],
                            "title": first_part,
                            "dates": second_part,
                            "description": "",
                            "bullets": []
                        }
                        current_subsection = new_role
                        if current_section:
                            current_section["subsections"].append(current_subsection)
                    else:
                        # Standalone job title - shouldn't happen in well-formed CV
                        pass

                else:
                    # This is a new company entry - look for job title in next few lines
                    company = first_part
                    location = second_part

                    # Look for job title, skipping empty lines and descriptions
                    title_found = False
                    for j in range(i + 1, min(i + 6, len(self.lines))):
                        check_line = self.lines[j].strip()

                        # Skip empty lines and company descriptions
                        if not check_line or (check_line.startswith('*') and not check_line.startswith('**')):
                            continue

                        # Found a job title line
                        if check_line.startswith('**') and '|' in check_line:
                            title_parts = check_line.split('|')
                            title, _ = self._extract_bold_text(title_parts[0])
                            dates = title_parts[1].strip() if len(title_parts) > 1 else ""

                            current_subsection = {
                                "company": company,
                                "location": location,
                                "title": title,
                                "dates": dates,
                                "description": "",
                                "bullets": []
                            }
                            if current_section:
                                current_section["subsections"].append(current_subsection)
                            processed_lines.add(j)  # Skip the title line when we get to it
                            title_found = True
                            break

                    # If no title found, still create the company entry (for cases with no title)
                    if not title_found:
                        current_subsection = {
                            "company": company,
                            "location": location,
                            "title": "",
                            "dates": "",
                            "description": "",
                            "bullets": []
                        }
                        if current_section:
                            current_section["subsections"].append(current_subsection)

            # Company descriptions (italic text after job entries)
            elif line.startswith('*') and not line.startswith('**') and current_subsection and not current_subsection["description"]:
                description = self._extract_italic_text(line)
                current_subsection["description"] = description

            # Bullet points
            elif line.startswith('• '):
                bullet_text = line[2:].strip()
                if current_subsection:
                    current_subsection["bullets"].append(bullet_text)
                elif current_section:
                    current_section["content"].append(("bullet", bullet_text))

            # Regular paragraphs (including bold text in non-Professional Experience sections)
            elif line and current_section and not line.startswith('#'):
                # Allow bold text in sections other than Professional Experience
                if (not any(line.startswith(x) for x in ['•', '|']) and
                    not (line.startswith('**') and current_section["title"] == "Professional Experience")):
                    current_section["content"].append(("paragraph", line))

        return data


class WordDocumentBuilder:
    """Build Word document with specified styling"""

    def __init__(self, style_config: StyleConfig):
        self.doc = Document()
        self.style_config = style_config
        self._setup_document()

    def _setup_document(self):
        """Set up document margins and default styles"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(self.style_config.styles["margins"]["top"])
            section.bottom_margin = Inches(self.style_config.styles["margins"]["bottom"])
            section.left_margin = Inches(self.style_config.styles["margins"]["left"])
            section.right_margin = Inches(self.style_config.styles["margins"]["right"])

    def add_header(self, name: str, title: str, contact: List[str]):
        """Add document header with name and contact info"""
        # Add name
        p = self.doc.add_paragraph()
        p.alignment = self.style_config.styles["header_align"]
        run = p.add_run(name)
        run.bold = True
        run.font.name = self.style_config.styles["font"]
        run.font.size = Pt(self.style_config.styles["name_size"])
        p.paragraph_format.space_after = Pt(self.style_config.styles["spacing"]["after_name"])

        # Add title
        if title:
            p = self.doc.add_paragraph()
            p.alignment = self.style_config.styles["header_align"]
            run = p.add_run(title)
            run.bold = True
            run.font.name = self.style_config.styles["font"]
            run.font.size = Pt(self.style_config.styles["contact_size"])
            p.paragraph_format.space_after = Pt(self.style_config.styles["spacing"]["after_title"])

        # Add contact info - handle both old pipe format and new separate lines format
        contact_items = []
        for item in contact:
            item = item.strip()

            # Check if already labeled (new format)
            if any(item.startswith(label) for label in ['Location:', 'Nationality:', 'Mobile:', 'Email:']):
                is_bold = False  # Match reference CV - contact fields not bold
                contact_items.append((item, is_bold))
            elif item.startswith('http') or 'linkedin' in item.lower():
                contact_items.append((item, False))  # LinkedIn not bold
            # Handle old pipe format (backwards compatibility)
            elif '@' in item and not item.startswith('Email:'):
                contact_items.append(("Email: " + item, False))
            elif (item.startswith('+') or any(char.isdigit() for char in item[:5])) and not item.startswith('Mobile:'):
                contact_items.append(("Mobile: " + item, False))
            elif 'australian' in item.lower() and 'citizen' in item.lower() and not item.startswith('Nationality:'):
                contact_items.append(("Nationality: " + item, False))
            elif ('perth' in item.lower() or 'australia' in item.lower()) and not item.startswith('Location:'):
                contact_items.append(("Location: " + item, False))
            else:
                contact_items.append((item, False))

        # Create separate paragraphs for each contact item
        for i, (contact_text, is_bold) in enumerate(contact_items):
            p = self.doc.add_paragraph()
            p.alignment = self.style_config.styles["header_align"]
            run = p.add_run(contact_text)
            run.bold = is_bold
            run.font.name = self.style_config.styles["font"]
            run.font.size = Pt(self.style_config.styles["contact_size"])

            # Explicitly set spacing - minimal for all except last item
            if i == len(contact_items) - 1:
                p.paragraph_format.space_after = Pt(self.style_config.styles["spacing"]["after_contact"])
            else:
                p.paragraph_format.space_after = Pt(0.1)  # Very small spacing to override defaults
                p.paragraph_format.space_before = Pt(0.1)

    def add_section(self, section: Dict):
        """Add a main section to the document"""
        # Check if this is the summary section
        is_summary = section["title"].lower() == "professional summary"

        # Section heading
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(self.style_config.styles["spacing"]["before_section"])

        # Removed keep settings to avoid excessive white space
        # if section["title"].lower() == "professional experience":
        #     p.paragraph_format.keep_with_next = True
        #     p.paragraph_format.keep_together = True

        run = p.add_run(section["title"])
        run.bold = False
        run.font.name = "Calibri Light"
        run.font.size = Pt(self.style_config.styles["h1_size"])
        p.paragraph_format.space_after = Pt(self.style_config.styles["spacing"]["after_section_title"])

        # Section content
        for i, (content_type, content) in enumerate(section["content"]):
            if content_type == "paragraph":
                p = self.doc.add_paragraph()
                p.alignment = self.style_config.styles["body_align"]
                # Use formatted text handling to process **bold** markers
                self._add_formatted_text(p, content, font_size=self.style_config.styles["body_size"])
                # Set line spacing
                p.paragraph_format.line_spacing = self.style_config.styles["line_spacing"]
                # Apply summary spacing if this is the summary section
                if is_summary:
                    p.paragraph_format.space_after = Pt(self.style_config.styles["spacing"]["after_summary"])

            elif content_type == "bullet":
                p = self.doc.add_paragraph(style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Add spacing between bullets
                if i < len(section["content"]) - 1:  # Not the last bullet
                    p.paragraph_format.space_after = Pt(self.style_config.styles["spacing"]["between_bullets"])
                # Handle bold text within bullets
                self._add_formatted_text(p, content, font_size=self.style_config.styles["bullet_size"])

        # Subsections (job entries)
        prev_company = None
        for i, subsection in enumerate(section["subsections"]):
            is_first_job = i == 0 and section["title"].lower() == "professional experience"
            # Check if this is the same company as the previous job
            is_same_company = (prev_company == subsection["company"] and
                             prev_company is not None)
            self._add_job_entry(subsection, is_first_job=is_first_job, skip_company=is_same_company)
            prev_company = subsection["company"]

    def _add_job_entry(self, job: Dict, is_first_job: bool = False, skip_company: bool = False):
        """Add a job entry with company, title, and bullets - with page break prevention"""
        # Company and location - only add if not skipping (i.e., not a duplicate)
        if not skip_company:
            company_p = self.doc.add_paragraph()
            company_p.paragraph_format.space_before = Pt(self.style_config.styles["spacing"]["before_job"])
            company_p.alignment = self.style_config.styles["body_align"]

            # Removed keep settings to avoid excessive white space
            # company_p.paragraph_format.keep_with_next = True
            # company_p.paragraph_format.keep_together = True

            # if is_first_job:
            #     company_p.paragraph_format.keep_together = True

            run = company_p.add_run(job["company"])
            run.bold = True
            run.font.name = self.style_config.styles["font"]
            run.font.size = Pt(self.style_config.styles["company_size"])

            if job["location"]:
                run = company_p.add_run(f" | {job['location']}")
                run.font.name = self.style_config.styles["font"]
                run.font.size = Pt(self.style_config.styles["company_size"])

            # Set spacing after company line to 0pt (was missing, causing 10pt default)
            company_p.paragraph_format.space_after = Pt(0)

        # Job title and dates - keep with next
        title_p = None
        if job["title"]:
            title_p = self.doc.add_paragraph()
            title_p.alignment = self.style_config.styles["body_align"]

            # Removed keep settings to avoid excessive white space
            # title_p.paragraph_format.keep_with_next = True
            # title_p.paragraph_format.keep_together = True

            run = title_p.add_run(job["title"])
            run.bold = True
            run.font.name = self.style_config.styles["font"]
            run.font.size = Pt(self.style_config.styles["role_size"])

            if job["dates"]:
                run = title_p.add_run(f" | {job['dates']}")
                run.font.name = self.style_config.styles["font"]
                run.font.size = Pt(self.style_config.styles["role_size"])

            title_p.paragraph_format.space_after = Pt(self.style_config.styles["spacing"]["after_job_title"])

        # Company description (italic text) - keep with first bullet
        desc_p = None
        if job.get("description"):
            desc_p = self.doc.add_paragraph()
            desc_p.alignment = self.style_config.styles["body_align"]
            desc_p.paragraph_format.line_spacing = self.style_config.styles["line_spacing"]

            # Removed keep settings to avoid excessive white space
            # desc_p.paragraph_format.keep_with_next = True
            # desc_p.paragraph_format.keep_together = True

            run = desc_p.add_run(job["description"])
            run.italic = True
            run.font.name = self.style_config.styles["font"]
            run.font.size = Pt(self.style_config.styles["description_size"])
            # Set color to 50% lighter (grey) - RGB(128, 128, 128)
            run.font.color.rgb = RGBColor(128, 128, 128)
            desc_p.paragraph_format.space_after = Pt(2)  # 2pt trailing space

        # Bullets with page break control
        for i, bullet in enumerate(job["bullets"]):
            bullet_p = self.doc.add_paragraph(style='List Bullet')
            bullet_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Removed keep settings to avoid excessive white space
            # if i < 2:
            #     bullet_p.paragraph_format.keep_with_next = True
            # bullet_p.paragraph_format.keep_together = True

            # Add spacing between bullets (except for the last one)
            if i < len(job["bullets"]) - 1:
                bullet_p.paragraph_format.space_after = Pt(self.style_config.styles["spacing"]["between_bullets"])

            self._add_formatted_text(bullet_p, bullet, font_size=self.style_config.styles["bullet_size"])

    def _add_formatted_text(self, paragraph, text: str, font_size: int = None):
        """Add text with bold formatting preserved"""
        if font_size is None:
            font_size = self.style_config.styles["body_size"]
        # Parse bold text marked with **
        parts = re.split(r'(\*\*[^*]+\*\*)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # Bold text - ensure it's properly formatted
                bold_content = part[2:-2]
                if bold_content:  # Don't add empty runs
                    run = paragraph.add_run(bold_content)
                    run.bold = True
                    run.font.name = self.style_config.styles["font"]
                    run.font.size = Pt(font_size)
            elif part:  # Don't add empty parts
                # Normal text
                run = paragraph.add_run(part)
                run.font.name = self.style_config.styles["font"]
                run.font.size = Pt(self.style_config.styles["body_size"])

    def save(self, output_path: Path):
        """Save the document to file"""
        self.doc.save(output_path)


def convert_markdown_to_docx(
    input_path: Path,
    output_path: Path,
    mode: str = "styled",
    template_path: Optional[Path] = None
) -> None:
    """
    Main conversion function

    Args:
        input_path: Path to markdown file
        output_path: Path for output Word document
        mode: "styled" or "ats" for different formatting
        template_path: Optional reference Word document for style extraction
    """
    # Read markdown content
    with open(input_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()

    # Parse markdown
    parser = MarkdownParser(markdown_content)
    parsed_data = parser.parsed_data

    # Create Word document with specified style
    style_config = StyleConfig(mode=mode)
    builder = WordDocumentBuilder(style_config)

    # Build document
    builder.add_header(
        parsed_data["name"],
        parsed_data["title"],
        parsed_data["contact"]
    )

    for section in parsed_data["sections"]:
        builder.add_section(section)

    # Save document
    builder.save(output_path)
    print(f"✓ Converted {input_path} to {output_path} (mode: {mode})")


def get_versioned_path(base_path: Path) -> Path:
    """
    Get next available version number for a file path.
    If file.docx exists, returns file_v2.docx, then file_v3.docx, etc.
    """
    if not base_path.exists():
        return base_path

    # Extract base name without extension
    stem = base_path.stem
    suffix = base_path.suffix
    parent = base_path.parent

    # Check if already has version suffix
    import re
    version_match = re.match(r'(.+)_v(\d+)$', stem)
    if version_match:
        base_name = version_match.group(1)
        current_version = int(version_match.group(2))
    else:
        base_name = stem
        current_version = 1

    # Find next available version
    version = current_version + 1
    while True:
        new_path = parent / f"{base_name}_v{version}{suffix}"
        if not new_path.exists():
            return new_path
        version += 1


def main():
    """CLI interface"""
    parser = argparse.ArgumentParser(
        description="Convert Markdown to Word with strict style control"
    )
    parser.add_argument("input", type=Path, help="Input markdown file")
    parser.add_argument(
        "-o", "--output",
        type=Path,
        help="Output Word file (default: input_name.docx)"
    )
    parser.add_argument(
        "-m", "--mode",
        choices=["styled", "ats", "readable"],
        default="styled",
        help="Output mode: 'styled' matches reference doc, 'ats' for ATS systems, 'readable' optimized for human reviewers"
    )
    parser.add_argument(
        "-t", "--template",
        type=Path,
        help="Reference Word document for style extraction"
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Overwrite existing file instead of creating versioned copy"
    )

    args = parser.parse_args()

    # Determine output path
    if args.output:
        output_path = args.output
    else:
        output_path = args.input.with_suffix('.docx')

    # Handle versioning unless overwrite flag is set
    if not args.overwrite and output_path.exists():
        output_path = get_versioned_path(output_path)
        if args.mode == "ats":
            output_path = args.input.with_name(
                args.input.stem + "_ats.docx"
            )
        elif args.mode == "readable":
            output_path = args.input.with_name(
                args.input.stem + "_readable.docx"
            )

    # Convert
    convert_markdown_to_docx(
        args.input,
        output_path,
        mode=args.mode,
        template_path=args.template
    )


if __name__ == "__main__":
    main()
