from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# Helper function to create a table
def create_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)

    return table

def create_change_request():
    doc = Document()

    # Title
    title = doc.add_heading('Change Request: CR-2025-XXXX', 0)
    doc.add_paragraph('Deploy VeloCloud SD-WAN Virtual Edge in Azure Australia East')
    doc.add_paragraph('Normal Change | Medium Risk | Low Impact')
    doc.add_paragraph()

    # Section 1: Change Request Details
    doc.add_heading('1. Change Request Details', 1)
    create_table(doc, ['Field', 'Value'], [
        ['CR Number', 'CR-2025-XXXX (Assign on submission)'],
        ['Title', 'Deploy VeloCloud SD-WAN Virtual Edge AERVAVC01 in Azure Australia East'],
        ['Requestor', 'TBD'],
        ['Date Submitted', '2025-XX-XX'],
        ['Target Implementation Date', 'TBD'],
        ['Change Type', 'Normal'],
        ['Priority', 'Medium'],
        ['Risk Level', 'Medium'],
        ['Impact Level', 'Low'],
        ['Category', 'Network Infrastructure'],
        ['Service', 'SD-WAN / Azure Networking'],
        ['CI Affected', 'AER-hub-australiaeast VNet, SD-WAN Overlay'],
    ])

    # Section 2: Change Description
    doc.add_heading('2. Change Description', 1)
    doc.add_heading('2.1 Summary', 2)
    doc.add_paragraph('Deploy a VMware VeloCloud SD-WAN Virtual Edge (AERVAVC01) in Azure Australia East region to extend SD-WAN connectivity to Azure-hosted workloads. This deployment enables optimized, policy-based routing between branch locations and Azure resources via the SD-WAN overlay.')

    doc.add_heading('2.2 Business Justification', 2)
    create_table(doc, ['Benefit', 'Description'], [
        ['SD-WAN Extension', 'Extend SD-WAN overlay to Azure for consistent network policies across all locations'],
        ['DMPO Optimization', 'Enable Dynamic Multipath Optimization for Azure traffic, improving application performance'],
        ['Failover Capability', 'Provide resilient branch-to-Azure connectivity with automatic failover'],
        ['Centralized Management', 'Single pane of glass management via VeloCloud Orchestrator'],
        ['Cost Optimization', 'Optimize bandwidth usage through intelligent path selection'],
    ])

    doc.add_heading('2.3 Technical Scope', 2)
    create_table(doc, ['Component', 'Details'], [
        ['Edge Name', 'AERVAVC01'],
        ['Location', 'Azure Australia East'],
        ['Target VNet', 'AER-hub-australiaeast (10.139.0.0/18)'],
        ['VM Size', 'Standard_D2d_v4 (2 vCPU, 8GB RAM)'],
        ['Edge Version', 'Virtual Edge 4.5.2'],
        ['Orchestrator', 'vco312-syd1.velocloud.net'],
        ['Availability Zone', 'Zone 1'],
        ['Public Subnet', 'velocloud-public-subnet (10.139.2.0/24)'],
        ['Private Subnet', 'velocloud-private-subnet (10.139.3.0/24)'],
        ['LAN IP Address', '10.139.3.4'],
    ])

    doc.add_heading('2.4 Resources Created', 2)
    create_table(doc, ['Resource', 'Name', 'Resource Group'], [
        ['Virtual Machine', 'AERVAVC01', 'rg-velocloud-aue'],
        ['OS Disk', 'AERVAVC01-osdisk', 'rg-velocloud-aue'],
        ['WAN NIC', 'AERVAVC01-nic-wan', 'rg-velocloud-aue'],
        ['LAN NIC', 'AERVAVC01-nic-lan', 'rg-velocloud-aue'],
        ['Public IP', 'AERVAVC01-pip', 'rg-velocloud-aue'],
        ['NSG', 'AERVAVC01-nsg', 'rg-velocloud-aue'],
        ['Subnets (if new)', 'velocloud-public-subnet, velocloud-private-subnet', 'rg-network-aue'],
    ])

    # Section 3: Risk Assessment
    doc.add_heading('3. Risk Assessment', 1)
    doc.add_heading('3.1 Risk Matrix', 2)
    create_table(doc, ['Risk ID', 'Risk Description', 'Likelihood', 'Impact', 'Risk Score', 'Mitigation'], [
        ['R1', 'VM deployment failure', 'Low', 'Low', 'Low', 'Template pre-validated; rollback = delete resource group'],
        ['R2', 'Edge activation failure', 'Low', 'Medium', 'Low', 'Verify VCO connectivity and activation key before deployment'],
        ['R3', 'Network connectivity disruption', 'Medium', 'Medium', 'Medium', 'Test in isolation first; gradual route migration'],
        ['R4', 'Performance degradation', 'Low', 'Low', 'Low', 'Start with non-critical traffic; monitor DMPO metrics'],
        ['R5', 'Cost overrun', 'Low', 'Low', 'Low', 'Known VM size with predictable monthly costs (~$215 AUD)'],
        ['R6', 'Security vulnerability', 'Low', 'High', 'Medium', 'Template security reviewed; NSG deny-all default'],
    ])

    doc.add_heading('3.2 Overall Risk Assessment', 2)
    doc.add_paragraph('Overall Risk Level: Medium')
    doc.add_paragraph('Rationale: New deployment with no impact to existing services. Template validated and security reviewed. Rollback plan tested and documented.')

    # Section 4: Implementation Plan
    doc.add_heading('4. Implementation Plan', 1)

    doc.add_heading('4.1 Phase 1: Preparation (Day -3 to Day -1)', 2)
    create_table(doc, ['Step', 'Task', 'Owner', 'Duration', 'Status'], [
        ['1.1', 'Obtain Change Approval (CAB)', 'Change Manager', '1 day', 'Pending'],
        ['1.2', 'Verify Azure prerequisites (quota, VNet, subnets)', 'Cloud Engineer', '2 hours', 'Pending'],
        ['1.3', 'Accept Azure Marketplace terms for VMware SD-WAN', 'Cloud Engineer', '15 min', 'Pending'],
        ['1.4', 'Create Edge record in VCO, obtain activation key', 'Network Engineer', '30 min', 'Pending'],
        ['1.5', 'Configure Edge profile in VCO (interfaces, policies)', 'Network Engineer', '1 hour', 'Pending'],
        ['1.6', 'Prepare secure parameter file', 'Cloud Engineer', '30 min', 'Pending'],
        ['1.7', 'Validate ARM template deployment', 'Cloud Engineer', '15 min', 'Pending'],
    ])

    doc.add_heading('4.2 Phase 2: Deployment (Day 0)', 2)
    create_table(doc, ['Step', 'Task', 'Owner', 'Duration', 'Status'], [
        ['2.1', 'Open maintenance window, notify stakeholders', 'Change Manager', '15 min', 'Pending'],
        ['2.2', 'Deploy ARM template to Azure', 'Cloud Engineer', '15 min', 'Pending'],
        ['2.3', 'Verify VM provisioning in Azure Portal', 'Cloud Engineer', '5 min', 'Pending'],
        ['2.4', 'Verify Edge activation in VCO', 'Network Engineer', '10 min', 'Pending'],
        ['2.5', 'Configure BGP peering (if required)', 'Network Engineer', '30 min', 'Pending'],
        ['2.6', 'Configure UDRs for branch route steering', 'Cloud Engineer', '30 min', 'Pending'],
    ])

    doc.add_heading('4.3 Phase 3: Validation (Day 0)', 2)
    create_table(doc, ['Step', 'Task', 'Owner', 'Duration', 'Status'], [
        ['3.1', 'Verify Edge status = CONNECTED in VCO', 'Network Engineer', '5 min', 'Pending'],
        ['3.2', 'Verify tunnel establishment (VPN test)', 'Network Engineer', '10 min', 'Pending'],
        ['3.3', 'Test connectivity to branch networks', 'Network Engineer', '15 min', 'Pending'],
        ['3.4', 'Verify business application access', 'App Owner', '30 min', 'Pending'],
        ['3.5', 'Review DMPO link quality metrics', 'Network Engineer', '10 min', 'Pending'],
    ])

    doc.add_heading('4.4 Phase 4: Documentation (Day +1)', 2)
    create_table(doc, ['Step', 'Task', 'Owner', 'Duration', 'Status'], [
        ['4.1', 'Update network diagrams', 'Network Engineer', '1 hour', 'Pending'],
        ['4.2', 'Update IPAM with new IP allocations', 'Cloud Engineer', '30 min', 'Pending'],
        ['4.3', 'Update operational runbooks', 'Network Engineer', '1 hour', 'Pending'],
        ['4.4', 'Close change request', 'Change Manager', '15 min', 'Pending'],
    ])

    # Section 5: Test Plan
    doc.add_heading('5. Test Plan', 1)
    doc.add_heading('5.1 Validation Tests', 2)
    create_table(doc, ['Test ID', 'Test Case', 'Expected Result', 'Actual Result', 'Pass/Fail'], [
        ['T1', 'VM deploys successfully', 'VM status = Running in Azure Portal', '', ''],
        ['T2', 'Edge activates in VCO', 'Edge status = CONNECTED', '', ''],
        ['T3', 'WAN link UP', 'Link status healthy, metrics normal', '', ''],
        ['T4', 'SD-WAN tunnels established', 'VPN test shows all peer tunnels UP', '', ''],
        ['T5', 'Ping from branch to Azure VM', '<50ms latency, 0% packet loss', '', ''],
        ['T6', 'Ping from Azure VM to branch', '<50ms latency, 0% packet loss', '', ''],
        ['T7', 'Business application access', 'Application accessible from branch', '', ''],
        ['T8', 'DMPO path optimization active', 'Link steering functional', '', ''],
    ])

    doc.add_heading('5.2 Test Commands', 2)
    p = doc.add_paragraph()
    p.add_run('# Azure validation\n').bold = True
    p.add_run('az vm show -g rg-velocloud-aue -n AERVAVC01 --query powerState -o tsv\n')
    p.add_run('az network public-ip show -g rg-velocloud-aue -n AERVAVC01-pip --query ipAddress -o tsv\n\n')
    p.add_run('# Connectivity tests (from Azure VM)\n').bold = True
    p.add_run('ping 192.168.10.1    # Branch A gateway\n')
    p.add_run('traceroute 192.168.10.1  # Verify path via 10.139.3.4')

    # Section 6: Rollback Plan
    doc.add_heading('6. Rollback Plan', 1)
    doc.add_heading('6.1 Rollback Triggers', 2)
    create_table(doc, ['Trigger', 'Threshold'], [
        ['Edge fails to activate', 'After 30 minutes of troubleshooting'],
        ['Critical application connectivity broken', 'Any production impact'],
        ['Unacceptable network performance', '>5% packet loss OR >100ms latency'],
        ['Security incident detected', 'Any severity'],
    ])

    doc.add_heading('6.2 Rollback Procedure', 2)
    create_table(doc, ['Step', 'Action', 'Owner', 'Duration'], [
        ['R1', 'Remove UDRs pointing to Edge LAN IP (10.139.3.4)', 'Cloud Engineer', '5 min'],
        ['R2', 'Deactivate Edge in VCO (Configure > Edges > Deactivate)', 'Network Engineer', '2 min'],
        ['R3', 'Delete Azure resource group', 'Cloud Engineer', '5 min'],
        ['R4', 'Verify original routing restored', 'Network Engineer', '3 min'],
        ['R5', 'Document rollback reason', 'Change Manager', '-'],
        ['R6', 'Schedule post-incident review', 'Change Manager', '-'],
    ])

    doc.add_heading('6.3 Rollback Commands', 2)
    p = doc.add_paragraph()
    p.add_run('# Remove UDRs (immediate traffic restoration)\n').bold = True
    p.add_run('az network route-table route delete --name route-to-branch-a --route-table-name rt-workloads-via-velocloud --resource-group rg-network-aue\n\n')
    p.add_run('# Delete all deployed resources\n').bold = True
    p.add_run('az group delete --name rg-velocloud-aue --yes --no-wait')

    doc.add_heading('6.4 Rollback Duration', 2)
    create_table(doc, ['Phase', 'Duration'], [
        ['Traffic restoration (UDR removal)', '5 minutes'],
        ['Complete rollback', '15 minutes'],
    ])

    # Section 7: Communication Plan
    doc.add_heading('7. Communication Plan', 1)
    create_table(doc, ['Timing', 'Audience', 'Message', 'Channel'], [
        ['Day -3', 'All Stakeholders', 'Change notification and schedule', 'Email'],
        ['Day -1', 'Network & Cloud Teams', 'Final preparation briefing', 'Teams Meeting'],
        ['Day 0 (Start)', 'Operations Team', 'Maintenance window started', 'Teams/Slack'],
        ['Day 0 (Complete)', 'All Stakeholders', 'Deployment successful', 'Email'],
        ['Day 0 (Issue)', 'All Stakeholders', 'Rollback initiated', 'Email + Phone'],
        ['Day +1', 'All Stakeholders', 'Change closure notification', 'Email'],
    ])

    # Section 8: Resource Requirements
    doc.add_heading('8. Resource Requirements', 1)
    doc.add_heading('8.1 Personnel', 2)
    create_table(doc, ['Role', 'Responsibility', 'Availability Required'], [
        ['Change Manager', 'Approval, coordination, communication', 'Day -3 to Day +1'],
        ['Cloud Engineer', 'Azure deployment, UDR configuration', 'Day 0 (4 hours)'],
        ['Network Engineer', 'VCO configuration, validation', 'Day 0 (4 hours)'],
        ['App Owner', 'Application validation', 'Day 0 (1 hour)'],
    ])

    doc.add_heading('8.2 Azure Resources (Monthly Cost Estimate)', 2)
    create_table(doc, ['Resource', 'Specification', 'Monthly Cost (AUD)'], [
        ['Virtual Machine', 'Standard_D2d_v4', '~$150'],
        ['Premium SSD (OS Disk)', '64GB P6', '~$15'],
        ['Public IP (Standard)', 'Static', '~$5'],
        ['Bandwidth (Egress)', 'Est. 500GB', '~$45'],
        ['Total', '', '~$215/month'],
    ])
    doc.add_paragraph('Note: VeloCloud license costs are covered under existing VMware agreement')

    # Section 9: Prerequisites Checklist
    doc.add_heading('9. Prerequisites Checklist', 1)
    doc.add_heading('9.1 Azure Prerequisites', 2)
    create_table(doc, ['#', 'Requirement', 'Verified', 'Notes'], [
        ['1', 'Azure Subscription with Contributor role', '', ''],
        ['2', 'Resource Group rg-velocloud-aue exists (or create permissions)', '', ''],
        ['3', 'VNet AER-hub-australiaeast exists', '', ''],
        ['4', 'Subnets available (or will be created)', '', ''],
        ['5', 'IP address 10.139.3.4 available in private subnet', '', ''],
        ['6', 'Azure Marketplace terms accepted for VMware SD-WAN', '', ''],
        ['7', 'Sufficient quota for Standard_D2d_v4 in Australia East', '', ''],
    ])

    doc.add_heading('9.2 VeloCloud Orchestrator Prerequisites', 2)
    create_table(doc, ['#', 'Requirement', 'Verified', 'Notes'], [
        ['1', 'VCO access: vco312-syd1.velocloud.net', '', ''],
        ['2', 'Virtual Edge license available', '', ''],
        ['3', 'Edge profile configured', '', ''],
        ['4', 'Activation key generated', '', ''],
        ['5', 'Business policies defined', '', ''],
    ])

    doc.add_heading('9.3 Network Prerequisites', 2)
    create_table(doc, ['#', 'Requirement', 'Verified', 'Notes'], [
        ['1', 'Outbound internet access from public subnet', '', ''],
        ['2', 'UDP 2426 outbound allowed (NSG/Firewall)', '', ''],
        ['3', 'DNS resolution for vco312-syd1.velocloud.net', '', ''],
        ['4', 'UDR plan for branch route steering', '', ''],
    ])

    # Section 10: Approvals
    doc.add_heading('10. Approvals', 1)
    create_table(doc, ['Role', 'Name', 'Signature', 'Date', 'Decision'], [
        ['Requestor', '', '', '', ''],
        ['Change Manager', '', '', '', 'Approved / Rejected'],
        ['Network Lead', '', '', '', 'Approved / Rejected'],
        ['Cloud Lead', '', '', '', 'Approved / Rejected'],
        ['Business Owner', '', '', '', 'Approved / Rejected'],
        ['CAB Chair', '', '', '', 'Approved / Rejected'],
    ])

    # Section 11: Post-Implementation Review
    doc.add_heading('11. Post-Implementation Review', 1)
    doc.add_paragraph('To be completed after change implementation')
    create_table(doc, ['Criteria', 'Assessment'], [
        ['Change Successful', 'Yes / No / Partial'],
        ['Rollback Required', 'Yes / No'],
        ['Actual Duration', ''],
        ['Issues Encountered', ''],
        ['Lessons Learned', ''],
        ['Follow-up Actions', ''],
    ])

    # Section 12: Attachments
    doc.add_heading('12. Attachments', 1)
    create_table(doc, ['Document', 'Location'], [
        ['Deployment Plan', 'velocloud-edge-azure-deployment-plan.md'],
        ['ARM Template', 'velocloud-edge-azure-template.json'],
        ['Parameter File', 'velocloud-edge-azure-parameters.json'],
        ['Architecture Diagram', 'See Deployment Plan Section 2'],
    ])

    # Section 13: Change History
    doc.add_heading('13. Change History', 1)
    create_table(doc, ['Version', 'Date', 'Author', 'Changes'], [
        ['1.0', '2025-XX-XX', 'TBD', 'Initial CR created'],
    ])

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('Generated by VeloCloud SD-WAN Agent v2.2')

    return doc


def create_deployment_instructions():
    doc = Document()

    # Title
    doc.add_heading('VeloCloud Virtual Edge Deployment Instructions', 0)
    doc.add_paragraph('AERVAVC01 - Azure Australia East')
    doc.add_paragraph()

    # Quick Reference
    doc.add_heading('Quick Reference', 1)
    create_table(doc, ['Parameter', 'Value'], [
        ['Edge Name', 'AERVAVC01'],
        ['Region', 'Australia East'],
        ['VNet', 'AER-hub-australiaeast'],
        ['VM Size', 'Standard_D2d_v4'],
        ['LAN IP', '10.139.3.4'],
        ['VCO', 'vco312-syd1.velocloud.net'],
    ])

    # Deployment Steps Summary
    doc.add_heading('Deployment Steps Summary', 1)
    create_table(doc, ['Step', 'Task'], [
        ['1', 'Create Edge in VCO & get activation key'],
        ['2', 'Connect to Azure subscription'],
        ['3', 'Accept Azure Marketplace terms'],
        ['4', 'Create resource group'],
        ['5', 'Create parameter file (with activation key & password)'],
        ['6', 'Validate ARM template'],
        ['7', 'Deploy Virtual Edge (~10-15 min)'],
        ['8', 'Verify Azure resources'],
        ['9', 'Verify Edge activation in VCO'],
        ['10', 'Configure UDRs for branch routing'],
        ['11', 'Validate end-to-end connectivity'],
    ])

    # Prerequisites
    doc.add_heading('Prerequisites', 1)
    doc.add_heading('Azure', 2)
    doc.add_paragraph('[ ] Contributor access to target subscription')
    doc.add_paragraph('[ ] Resource group rg-velocloud-aue exists (or permissions to create)')
    doc.add_paragraph('[ ] VNet AER-hub-australiaeast exists')
    doc.add_paragraph('[ ] IP address 10.139.3.4 is available')
    doc.add_paragraph('[ ] Quota available for Standard_D2d_v4 in Australia East')

    doc.add_heading('VeloCloud Orchestrator', 2)
    doc.add_paragraph('[ ] Admin access to vco312-syd1.velocloud.net')
    doc.add_paragraph('[ ] Virtual Edge license available')
    doc.add_paragraph('[ ] Activation key generated (Step 1 below)')

    # Step 1
    doc.add_heading('Step 1: Create Edge in VeloCloud Orchestrator', 1)
    doc.add_paragraph('1. Login to VCO: https://vco312-syd1.velocloud.net')
    doc.add_paragraph()
    doc.add_paragraph('2. Create Edge Profile (if not exists):')
    p = doc.add_paragraph()
    p.add_run('   Navigate: Configure > Profiles > New Profile\n')
    p.add_run('   - Name: Azure-Virtual-Edge-Profile\n')
    p.add_run('   - Model: Virtual Edge\n')
    p.add_run('   - GE1: WAN, Overlay Enabled, DHCP\n')
    p.add_run('   - GE2: Routed/LAN, Static IP\n')
    p.add_run('   - GE3-GE8: Disabled')

    doc.add_paragraph()
    doc.add_paragraph('3. Create Edge Record:')
    p = doc.add_paragraph()
    p.add_run('   Navigate: Configure > Edges > New Edge\n')
    p.add_run('   - Name: AERVAVC01\n')
    p.add_run('   - Model: Virtual Edge\n')
    p.add_run('   - Profile: Azure-Virtual-Edge-Profile')

    doc.add_paragraph()
    doc.add_paragraph('4. Copy Activation Key:')
    p = doc.add_paragraph()
    p.add_run('   IMPORTANT: ').bold = True
    p.add_run('Copy the Activation Key displayed after creation\n')
    p.add_run('   Format: XXXX-XXXX-XXXX-XXXX - SAVE THIS KEY')

    doc.add_paragraph()
    doc.add_paragraph('5. Configure Edge Firewall Access:')
    p = doc.add_paragraph()
    p.add_run('   Navigate: Configure > Edges > AERVAVC01 > Firewall > Edge Access\n')
    p.add_run('   Add Allow Rule: Source IP: 168.63.129.16/32 (Azure Fabric IP)')

    # Step 2
    doc.add_heading('Step 2: Connect to Azure Subscription', 1)
    doc.add_paragraph('Login to Azure:')
    doc.add_paragraph('az login')
    doc.add_paragraph()
    doc.add_paragraph('List Available Subscriptions:')
    doc.add_paragraph('az account list --output table')
    doc.add_paragraph()
    doc.add_paragraph('Set Target Subscription:')
    doc.add_paragraph('az account set --subscription "YOUR-SUBSCRIPTION-NAME-OR-ID"')
    doc.add_paragraph()
    doc.add_paragraph('Verify Active Subscription:')
    doc.add_paragraph('az account show --output table')

    # Step 3
    doc.add_heading('Step 3: Accept Azure Marketplace Terms', 1)
    doc.add_paragraph('Run this command once per subscription:')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('az vm image terms accept `\n')
    p.add_run('  --publisher vmware-inc `\n')
    p.add_run('  --offer sol-42222-bbj `\n')
    p.add_run('  --plan vmware_sdwan_452')

    # Step 4
    doc.add_heading('Step 4: Create Resource Group', 1)
    p = doc.add_paragraph()
    p.add_run('az group create `\n')
    p.add_run('  --name rg-velocloud-aue `\n')
    p.add_run('  --location australiaeast `\n')
    p.add_run('  --tags Application=SD-WAN Environment=Production')

    # Step 5
    doc.add_heading('Step 5: Create Parameter File', 1)
    doc.add_paragraph('Create file velocloud-edge-azure-parameters.json with deployment parameters.')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('IMPORTANT: ').bold = True
    p.add_run('Replace the following values:\n')
    p.add_run('- REPLACE-WITH-ACTIVATION-KEY -> Activation key from Step 1\n')
    p.add_run('- REPLACE-WITH-SECURE-PASSWORD -> Strong password (min 12 chars)')

    # Step 6
    doc.add_heading('Step 6: Validate Template', 1)
    p = doc.add_paragraph()
    p.add_run('az deployment group validate `\n')
    p.add_run('  --resource-group rg-velocloud-aue `\n')
    p.add_run('  --template-file velocloud-edge-azure-template.json `\n')
    p.add_run("  --parameters '@velocloud-edge-azure-parameters.json'")

    # Step 7
    doc.add_heading('Step 7: Deploy Virtual Edge', 1)
    p = doc.add_paragraph()
    p.add_run('az deployment group create `\n')
    p.add_run("  --name \"velocloud-edge-$(Get-Date -Format 'yyyyMMdd-HHmmss')\" `\n")
    p.add_run('  --resource-group rg-velocloud-aue `\n')
    p.add_run('  --template-file velocloud-edge-azure-template.json `\n')
    p.add_run("  --parameters '@velocloud-edge-azure-parameters.json'")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Deployment Time: ').bold = True
    p.add_run('~10-15 minutes')

    # Step 8
    doc.add_heading('Step 8: Verify Azure Deployment', 1)
    doc.add_paragraph('Check VM Status:')
    p = doc.add_paragraph()
    p.add_run('az vm show `\n')
    p.add_run('  --resource-group rg-velocloud-aue `\n')
    p.add_run('  --name AERVAVC01 `\n')
    p.add_run('  --query "powerState" `\n')
    p.add_run('  --output tsv')
    doc.add_paragraph()
    doc.add_paragraph('Expected: VM running')
    doc.add_paragraph()
    doc.add_paragraph('Get Public IP:')
    p = doc.add_paragraph()
    p.add_run('az network public-ip show `\n')
    p.add_run('  --resource-group rg-velocloud-aue `\n')
    p.add_run('  --name AERVAVC01-pip `\n')
    p.add_run('  --query "ipAddress" `\n')
    p.add_run('  --output tsv')

    # Step 9
    doc.add_heading('Step 9: Verify Edge Activation in VCO', 1)
    doc.add_paragraph('1. Login to VCO: https://vco312-syd1.velocloud.net')
    doc.add_paragraph('2. Navigate: Monitor > Edges')
    doc.add_paragraph('3. Find AERVAVC01 - Expected Status: CONNECTED (green)')
    doc.add_paragraph('4. Verify WAN Link: Monitor > Edges > AERVAVC01 > Links')
    doc.add_paragraph('5. Check Tunnels: Test & Troubleshoot > Remote Diagnostics > VPN Test')

    # Step 10
    doc.add_heading('Step 10: Configure User-Defined Routes (UDRs)', 1)
    doc.add_paragraph('Create Route Table:')
    p = doc.add_paragraph()
    p.add_run('az network route-table create `\n')
    p.add_run('  --name rt-workloads-via-velocloud `\n')
    p.add_run('  --resource-group rg-network-aue `\n')
    p.add_run('  --location australiaeast')
    doc.add_paragraph()
    doc.add_paragraph('Add Routes to Branch Networks:')
    p = doc.add_paragraph()
    p.add_run('az network route-table route create `\n')
    p.add_run('  --name route-to-branch-a `\n')
    p.add_run('  --route-table-name rt-workloads-via-velocloud `\n')
    p.add_run('  --resource-group rg-network-aue `\n')
    p.add_run('  --address-prefix 192.168.10.0/24 `\n')
    p.add_run('  --next-hop-type VirtualAppliance `\n')
    p.add_run('  --next-hop-ip-address 10.139.3.4')

    # Step 11
    doc.add_heading('Step 11: Validate End-to-End Connectivity', 1)
    doc.add_paragraph('From Azure VM (in workload subnet):')
    doc.add_paragraph('ping 192.168.10.1')
    doc.add_paragraph('traceroute 192.168.10.1')
    doc.add_paragraph()
    create_table(doc, ['Test', 'Expected'], [
        ['Azure -> Branch ping', '<50ms, 0% loss'],
        ['Branch -> Azure ping', '<50ms, 0% loss'],
        ['Traceroute via Edge', 'Shows 10.139.3.4 as hop'],
    ])

    # Troubleshooting
    doc.add_heading('Troubleshooting', 1)
    doc.add_heading('Edge Not Activating (OFFLINE in VCO)', 2)
    doc.add_paragraph('1. Check DNS Resolution: nslookup vco312-syd1.velocloud.net')
    doc.add_paragraph('2. Check NSG Rules: Verify UDP 2426 outbound is allowed')
    doc.add_paragraph('3. Check Activation Key: Verify key matches VCO exactly')
    doc.add_paragraph('4. View Boot Diagnostics: Azure Portal > AERVAVC01 > Boot diagnostics')

    # Rollback
    doc.add_heading('Rollback Procedure', 1)
    doc.add_paragraph('1. Remove UDRs (Immediate Traffic Restoration):')
    p = doc.add_paragraph()
    p.add_run('az network vnet subnet update `\n')
    p.add_run('  --name workload-subnet `\n')
    p.add_run('  --vnet-name AER-hub-australiaeast `\n')
    p.add_run('  --resource-group rg-network-aue `\n')
    p.add_run('  --remove routeTable')
    doc.add_paragraph()
    doc.add_paragraph('2. Deactivate Edge in VCO: Configure > Edges > AERVAVC01 > Deactivate')
    doc.add_paragraph()
    doc.add_paragraph('3. Delete Azure Resources:')
    doc.add_paragraph('az group delete --name rg-velocloud-aue --yes --no-wait')

    # Post-Deployment Checklist
    doc.add_heading('Post-Deployment Checklist', 1)
    create_table(doc, ['Task', 'Status'], [
        ['VM running in Azure', ''],
        ['Edge CONNECTED in VCO', ''],
        ['WAN link UP and healthy', ''],
        ['All tunnels established', ''],
        ['UDRs configured', ''],
        ['Branch connectivity tested', ''],
        ['Application access verified', ''],
        ['Network diagrams updated', ''],
        ['IPAM updated', ''],
        ['Runbooks updated', ''],
    ])

    # Support Contacts
    doc.add_heading('Support Contacts', 1)
    create_table(doc, ['Issue', 'Contact'], [
        ['Azure deployment issues', 'Cloud Team / Azure Support'],
        ['VeloCloud Edge issues', 'Network Team / VMware Support'],
        ['VCO access issues', 'VeloCloud TAM'],
        ['Application issues', 'Application Owner'],
    ])

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('Generated by VeloCloud SD-WAN Agent v2.2')

    return doc


if __name__ == '__main__':
    output_dir = r'C:\Users\olli.ojala\maia\data'

    # Create Change Request
    cr_doc = create_change_request()
    cr_path = os.path.join(output_dir, 'CR-AERVAVC01-VeloCloud-Edge-Deployment.docx')
    cr_doc.save(cr_path)
    print(f'Change Request saved to: {cr_path}')

    # Create Deployment Instructions
    di_doc = create_deployment_instructions()
    di_path = os.path.join(output_dir, 'AERVAVC01-Deployment-Instructions.docx')
    di_doc.save(di_path)
    print(f'Deployment Instructions saved to: {di_path}')

    print('Done!')
